"""
Word报告写入器

处理电池分析结果的Word（docx）文件写入
"""

import os
import math
import logging
import datetime
import importlib.resources
from pathlib import Path

from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx import Document

from battery_analysis.utils.writers import word_utils
from battery_analysis.utils import numeric_utils
from battery_analysis.utils.exceptions import BatteryAnalysisException
from battery_analysis.utils.report_coordinator import compute_report_content_base, match_battery_type
from battery_analysis.utils.writers.statistics_utils import (
    compute_list_cpt, compute_statistics,
)
from battery_analysis import __version__


logger = logging.getLogger(__name__)


class WordReportWriter:
    """Word报告写入器"""

    def __init__(self, strResultPath: str, listTestInfo: list, listBatteryInfo: list,
                 equipment_info: dict | None = None) -> None:
        self._equipment_info = equipment_info or {}
        self.strResultPath = strResultPath
        self.listTestInfo = listTestInfo
        self.listBatteryInfo = listBatteryInfo

        # 计算电流/电压等级信息
        self.listCurrentLevel = listTestInfo[14]
        self.listVoltageLevel = listTestInfo[15]
        self.intCurrentLevelNum = len(self.listCurrentLevel)
        self.intVoltageLevelNum = len(self.listVoltageLevel)

        # 计算文件电流类型字符串
        self.strFileCurrentType = ""
        for c in range(self.intCurrentLevelNum):
            self.strFileCurrentType += f"{self.listCurrentLevel[c]}-"
        self.strFileCurrentType = self.strFileCurrentType[:-1]

        # 电池信息
        self.listBatteryCharge = self.listBatteryInfo[0]
        self.listBatteryName = self.listBatteryInfo[1]
        self.intBatteryNum = len(self.listBatteryName)

        # 从strResultPath提取td（日期部分）
        basename = os.path.basename(strResultPath)
        if "_v" in basename:
            td = basename.split("_v")[0]
        else:
            td = "00000000"

        # 图像路径
        self.listPngPath = []
        for b in range(self.intCurrentLevelNum):
            self.listPngPath.append(
                (f"{self.strResultPath}/Image_UseableCapacityOver"
                 f"CutoffVoltage{self.listCurrentLevel[b]}mALoad.png"))
        self.strFilteredPngPath = (
            f"{self.strResultPath}/Image_FilteredLoadVoltageOverCharge.png"
        )

        # Excel路径（用于相对链接）
        safe_temperature = self.listTestInfo[7].replace(':', '_')
        self.strResultXlsxPath = (
            f"{self.strResultPath}/{self.listTestInfo[4]}_{self.listTestInfo[2]}_{self.listTestInfo[3]}"
            f"_{self.strFileCurrentType}_{safe_temperature}.xlsx"
        )
        self.strSampleXlsxPath = (
            f"{self.strResultPath}/Sample_{self.listTestInfo[4]}_{self.listTestInfo[2]}_{self.listTestInfo[3]}"
            f"_{self.strFileCurrentType}_{safe_temperature}.xlsx"
        )

        # Word模板路径
        if self.intCurrentLevelNum <= 4:
            template_filename = "Battery Measurement Report of TypeC TypeA_TypeD.docx"
        else:
            template_filename = "Battery Measurement Report of TypeC TypeA_TypeD_MP.docx"

        try:
            pkg_template = (
                importlib.resources.files('battery_analysis')
                / 'templates'
                / template_filename
            )
            if pkg_template.is_file():
                self.strSampleReportWordPath = str(pkg_template)
                logging.info("从包内模板目录加载Word模板: %s", self.strSampleReportWordPath)
            else:
                raise FileNotFoundError
        except (TypeError, ModuleNotFoundError, FileNotFoundError):
            self.strSampleReportWordPath = str(
                Path(self.strResultPath) / f"../../0_doc/{template_filename}"
            )
            logging.info("从外部0_doc目录加载Word模板: %s", self.strSampleReportWordPath)

        # Word输出路径
        report_name = (
            f"{self.listTestInfo[4]}_{self.listTestInfo[2]}_DC{self.listTestInfo[5]}"
            f"_TD{td}_v{self.listTestInfo[16]}.docx"
        )
        result_dir = Path(self.strResultPath).parent
        self.strReportWordPath = str(result_dir / report_name)

        # 文本替换占位符
        self.listTextToReplace = [
            "TypeA", "TypeB", "TypeC", "TypeD",
            "TypeE", "TypeF", "TypeG", "StrA",
            "StrB", "StrC", "StrD", "StrF"
        ]
        self.listImageToReplace = ["<<Image_FilteredLoadVoltageOverCharge>>"]
        for i in range(10):
            self.listImageToReplace.append(
                f"<<Image_UseableCapacityOverCutoffVoltage{i}>>")
            self.listImageToReplace.append(
                f"<<Title_UseableCapacityOverCutoffVoltage{i}>>")

        # 电池类型匹配
        strBatteryType = match_battery_type(self.listTestInfo[2])

        # strStrF（颜色/电流等级字符串）
        self.listColorName = ["red = ", "blue = ", "yellow = ",
                              "violet = ", "green = ", "orange = ", "black1 = ", "black2 = "]
        strStrF = ""
        for c in range(self.intCurrentLevelNum):
            strStrF += f"{self.listColorName[c]}{self.listCurrentLevel[c]}mA, "
        strStrF = strStrF[:-2]

        self.listTestInfoForReplace = [
            self.listTestInfo[2], self.listTestInfo[3], self.listTestInfo[4],
            self.listTestInfo[5], self.listTestInfo[7], self.listTestInfo[11],
            strBatteryType, None, None, None, None, strStrF
        ]

    def _get_equip_value(self, dotted_key: str, fallback: str = "") -> str:
        """从 _equipment_info 字典读取带点号的键"""
        parts = dotted_key.split(".")
        value = self._equipment_info
        try:
            for p in parts:
                value = value[p]
            if isinstance(value, str):
                return value
            if isinstance(value, (list, tuple)):
                return ", ".join(str(v) for v in value)
        except (KeyError, TypeError, IndexError):
            pass
        return fallback

    # ── 写入Version History表 ──

    def _write_version_history(self, doc):
        """在Word文档中写入版本历史表"""
        table = doc.add_table(9, 4, style='Grid Table 4 Accent 3')
        for r in range(9):
            table.rows[r].height = Cm(0.6)
            for c in range(4):
                if c < 4:
                    table.cell(0, c).width = Cm(2.8)
                cell = table.cell(r, c)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        headers = ["Date", "Version", "Editor", "Changes"]
        for j, header in enumerate(headers):
            table.cell(0, j).paragraphs[0].add_run(header).font.size = Pt(10)

        version_data = [
            (0, datetime.datetime.now().strftime("%Y.%m.%d"), {"bold": False}),
            (1, "1.0", {}),
            (2, self.listTestInfo[18] if len(self.listTestInfo) > 18 else "", {}),
            (3, "Initial version", {}),
        ]
        for col, content, properties in version_data:
            text = table.cell(1, col).paragraphs[0].add_run(content)
            text.font.size = Pt(10)
            for prop, value in properties.items():
                setattr(text, prop, value)

        return table

    # ── 写入Test Information表 ──

    def _write_test_information(self, doc):
        """在Word文档中写入测试信息表"""
        table = doc.add_table(5, 2, style='Table Grid')
        table.cell(0, 1).width = Cm(10)
        for r in range(5):
            for c in range(2):
                cell = table.cell(r, c)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
                cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)

        test_info_headers = [
            "Test Equipment",
            "Software Versions",
            "middle Machines",
            "Test Units",
            "Data Processing Platforms",
        ]
        for r, header in enumerate(test_info_headers):
            table.cell(r, 0).paragraphs[0].add_run(header).font.size = Pt(10)

        test_info_data = [
            (0, lambda: self._equipment_info.get("testEquipment", "")),
            (1, lambda: (
                f"BTS Server Version: {self._get_equip_value('softwareVersions.btsServer')}\n"
                f"BTS Client Version: {self._get_equip_value('softwareVersions.btsClient')}\n"
                f"BTSDA (Data Analysis) Version: {self._get_equip_value('softwareVersions.btsda')}"
            )),
            (2, lambda: (
                f"Model: {self._get_equip_value('middleMachines.model')}\n"
                f"Hardware Version: {self._get_equip_value('middleMachines.hardwareVersion')}\n"
                f"Serial Number: {self._get_equip_value('middleMachines.serialNumber')}\n"
                f"Firmware Version: {self._get_equip_value('middleMachines.firmwareVersion')}\n"
                f"Device Type: {self._get_equip_value('middleMachines.deviceType')}"
            )),
            (3, lambda: (
                f"Model: {self._get_equip_value('testUnits.model')}\n"
                f"Hardware Version: {self._get_equip_value('testUnits.hardwareVersion')}\n"
                f"Firmware Version: {self._get_equip_value('testUnits.firmwareVersion')}"
            )),
            (4, lambda: f"Battery Analyzer-v{__version__}"),
        ]
        for row, content_func in test_info_data:
            content = content_func()
            table.cell(row, 1).paragraphs[0].add_run(f"{content}").font.size = Pt(10)

        return table

    # ── 写入Statistical Results表 ──

    def _write_statistical_results(self, doc, stats):
        """在Word文档中写入统计结果表"""
        table = doc.add_table(
            self.intCurrentLevelNum + 1, self.intVoltageLevelNum + 1, style='Table Grid')
        for c in range(self.intCurrentLevelNum + 1):
            for v in range(self.intVoltageLevelNum + 1):
                cell = table.cell(c, v)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                if c == 0 and v == 0:
                    table.rows[c].height = Cm(0.7)
                    table.cell(c, v).width = Cm(3.55)
                    text = cell.paragraphs[0].add_run("Statisticals\nResults")
                    text.font.size = Pt(12)
                    text.bold = True
                    word_utils.table_set_bg_color(cell, '#BFBFBF')
                elif c == 0 and v > 0:
                    table.cell(c, v).width = Cm(3.55)
                    text1 = cell.paragraphs[0].add_run(f"Cut-off Voltage\n")
                    text1.font.size = Pt(12)
                    text2 = cell.paragraphs[0].add_run(f"{self.listVoltageLevel[v - 1]}V")
                    text2.font.bold = True
                    text2.font.size = Pt(12)
                    word_utils.table_set_bg_color(cell, '#F2F2F2')
                elif c > 0 and v == 0:
                    table.rows[c].height = Cm(2.35)
                    text1 = cell.paragraphs[0].add_run(f"Pulse Current\n")
                    text1.font.size = Pt(12)
                    text2 = cell.paragraphs[0].add_run(f"{self.listCurrentLevel[c - 1]}mA")
                    text2.font.bold = True
                    text2.font.size = Pt(12)
                    word_utils.table_set_bg_color(cell, '#F2F2F2')
                else:
                    text = cell.paragraphs[0].add_run((
                        f"μ: {round(stats['mean'][c - 1][v - 1])}mAh\n"
                        f"Median: {round(stats['med'][c - 1][v - 1])}mAh\n"
                        f"σ: {round(stats['std'][c - 1][v - 1])}mAh\n"
                        f"μ - 3σ: {round(stats['mm3s'][c - 1][v - 1])}mAh\n"
                        f"μ - 2σ: {round(stats['mm2s'][c - 1][v - 1])}mAh\n"
                        f"μ + 2σ: {round(stats['mp2s'][c - 1][v - 1])}mAh\n"
                        f"μ + 3σ: {round(stats['mp3s'][c - 1][v - 1])}mAh\n"
                        f"Minimum: {round(stats['min'][c - 1][v - 1])}mAh\n"
                        f"Maximum: {round(stats['max'][c - 1][v - 1])}mAh"))
                    text.font.size = Pt(7)
                    cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.EXACTLY
                    cell.paragraphs[0].paragraph_format.line_spacing = Pt(10)

        return table

    # ── 计算Overview表内容 ──

    def _prepare_overview_content(self, doc, stats):
        """计算Overview表所需的位置参数和内容列表"""
        base = compute_report_content_base(
            self.listCurrentLevel, self.intCurrentLevelNum,
            self.listVoltageLevel, self.intVoltageLevelNum,
            self.listTestInfo, self.listBatteryInfo,
            self.strSampleXlsxPath, self.strResultXlsxPath, self.strReportWordPath,
            stats)

        intTestDateStartRow = base['intTestProfileStartLine'] + 11

        # Word 特有的 listTestInfoForReplace 副作用
        if base['strResult'] == "Pass":
            self.listTestInfoForReplace[8] = "meets"
            self.listTestInfoForReplace[10] = "Pass"
        else:
            self.listTestInfoForReplace[8] = "doesn't meet"
            self.listTestInfoForReplace[10] = "Fail"
        self.listTestInfoForReplace[9] = base['strRequiredUseableCapacityPercentage']
        self.listTestInfoForReplace[7] = (
            f"{math.floor(100 * stats['mm2s'][base['intPosiMaxmA']][base['intPosi2V25']] / int(self.listTestInfo[9]))}%")

        return {
            **base,
            'intTestDateStartRow': intTestDateStartRow,
        }

    # ── 写入Overview表（Word） ──

    def _write_overview_table(self, doc, content, stats):
        """在Word文档中写入Overview表"""
        intPosiMaxmA = content['intPosiMaxmA']
        intPosi2V25 = content['intPosi2V25']
        intTestProfileStartLine = content['intTestProfileStartLine']
        intTestDateStartRow = content['intTestDateStartRow']
        intActualMeasuredCapacityLength = content['intActualMeasuredCapacityLength']
        listStrItems = content['listStrItems']
        listStrContent = content['listStrContent']

        table = doc.add_table(
            intTestDateStartRow + 6, 1 + self.intVoltageLevelNum * 2, style='Table Grid')
        for row in range(intTestDateStartRow + 6):
            table.rows[row].height = Cm(0.5)
            for col in range(1 + self.intVoltageLevelNum * 2):
                cell = table.cell(row, col)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                if row == intTestDateStartRow + 3:
                    word_utils.table_set_bg_color(cell, '#FFFF00')

        # merge cells
        for row in range(intTestDateStartRow + 6):
            if row <= intTestProfileStartLine + 6 or row >= intTestDateStartRow:
                cell1 = table.cell(row, 1)
                for col in range(1, self.intVoltageLevelNum * 2):
                    cell2 = table.cell(row, col + 1)
                    cell1.merge(cell2)
            else:
                if row == intTestProfileStartLine + 7:
                    cell1 = table.cell(row, 1)
                    for col in range(1, self.intVoltageLevelNum):
                        cell2 = table.cell(row, col + 1)
                        cell1.merge(cell2)
                    cell1 = table.cell(row, 1 + self.intVoltageLevelNum)
                    for col in range(1 + self.intVoltageLevelNum, self.intVoltageLevelNum * 2):
                        cell2 = table.cell(row, col + 1)
                        cell1.merge(cell2)
                else:
                    for v in range(self.intVoltageLevelNum):
                        cell1 = table.cell(row, 1 + v * 2)
                        cell2 = table.cell(row, 2 + v * 2)
                        cell1.merge(cell2)

        cell1 = table.cell(intTestProfileStartLine + 8, 0)
        cell2 = table.cell(intTestProfileStartLine + 9, 0)
        cell1.merge(cell2)
        cell2 = table.cell(intTestProfileStartLine + 10, 0)
        cell1.merge(cell2)

        # 写入标签列
        for i in range(4):
            table.cell(i, 0).paragraphs[0].add_run(listStrItems[i])
        if intTestProfileStartLine == 4:
            table.cell(3, 0).paragraphs[0].text = table.cell(
                3, 0).paragraphs[0].text.replace(listStrItems[3], "")
        for i in range(4, 12):
            table.cell(intTestProfileStartLine + (i - 4),
                       0).paragraphs[0].add_run(listStrItems[i])
        table.cell(intTestProfileStartLine + 8,
                   0).paragraphs[0].add_run(listStrItems[13])
        for i in range(14, 20):
            table.cell(intTestDateStartRow + (i - 14),
                       0).paragraphs[0].add_run(listStrItems[i])

        # 写入内容列
        for i in range(4):
            table.cell(i, 1).paragraphs[0].add_run(listStrContent[i])
        if intTestProfileStartLine == 4:
            table.cell(3, 1).paragraphs[0].text = table.cell(
                3, 1).paragraphs[0].text.replace(listStrContent[3], "")
        table.cell(intTestProfileStartLine, 1).paragraphs[0].text = ""
        if len(listStrContent[4].split("\\")) == 1:
            table.cell(intTestProfileStartLine,
                       1).paragraphs[0].add_run(listStrContent[4])
        else:
            word_utils.add_hyperlink(table.cell(
                intTestProfileStartLine, 1).paragraphs[0], listStrContent[4],
                listStrContent[4].split("\\")[-1])
        for i in range(5, 12):
            table.cell(intTestProfileStartLine + (i - 4),
                       1).paragraphs[0].add_run(listStrContent[i])
        table.cell(intTestProfileStartLine + 7, 1 +
                   self.intVoltageLevelNum).paragraphs[0].add_run(listStrContent[12])

        for v in range(self.intVoltageLevelNum):
            text1 = table.cell(
                intTestProfileStartLine + 8, 1 + v * 2).paragraphs[0].add_run(f"{self.listVoltageLevel[v]}V")
            text2 = table.cell(intTestProfileStartLine + 9, 1 + v * 2).paragraphs[0].add_run(
                f"{round(stats['mm2s'][intPosiMaxmA][v], 2)}")
            text3 = table.cell(intTestProfileStartLine + 10, 1 + v * 2).paragraphs[0].add_run(
                f"{math.floor(100 * stats['mm2s'][intPosiMaxmA][v] / int(listStrContent[10]))}%")
            if v == intPosi2V25:
                text1.font.bold = True
                text2.font.bold = True
                text3.font.bold = True

        for i in range(14, 20):
            if i == 18:
                table.cell(intTestDateStartRow + (i - 14), 1).paragraphs[0].text = ""
                word_utils.add_hyperlink(table.cell(intTestDateStartRow + (i - 14), 1).paragraphs[0],
                                         listStrContent[i], listStrContent[i].split("\\")[-1])
            else:
                table.cell(intTestDateStartRow + (i - 14),
                           1).paragraphs[0].add_run(listStrContent[i])

        for row in range(intTestDateStartRow + 6):
            for col in range(1 + self.intVoltageLevelNum * 2):
                cell = table.cell(row, col)
                runs = cell.paragraphs[0].runs
                for run in runs:
                    run.font.size = Pt(9)
        table.cell(0, 0).width = Cm(27)

        return table

    # ── 文本替换、图片插入、表格插入 ──

    def _replace_and_insert(self, doc, tables):
        """替换占位符文本、插入图片，并在指定位置插入表格"""
        tableOverview = tables['overview']
        tableVersionHistory = tables['version_history']
        tableTestInformation = tables['test_information']
        tableStatisticalsResults = tables['statistical_results']

        bInsertOverview = False
        bInsertVersionHistory = False
        bInsertTestInformation = False
        bInsertStatisticalsResults = False
        intStepOut = 0

        for paragraph in doc.paragraphs:
            modified = False

            # 替换文本
            for t in range(len(self.listTextToReplace)):
                if self.listTextToReplace[t] in paragraph.text:
                    modified = True
                    if self.listTextToReplace[t] == "StrD":
                        paragraph.text = paragraph.text.replace(
                            self.listTextToReplace[t], "")
                        text = paragraph.add_run(f"{self.listTestInfoForReplace[t]}")
                        text.font.bold = True
                        paragraph.add_run(".")
                    else:
                        paragraph.text = paragraph.text.replace(
                            self.listTextToReplace[t], f"{self.listTestInfoForReplace[t]}")

            # 替换图片
            if not modified:
                for i in range(len(self.listImageToReplace)):
                    if self.listImageToReplace[i] in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            self.listImageToReplace[i], "")
                        if i < 2 * len(self.listPngPath) + 1:
                            if i == 0:
                                paragraph.add_run("").add_picture(
                                    self.strFilteredPngPath, width=Cm(15))
                            elif i % 2 == 1:
                                paragraph.add_run("").add_picture(
                                    self.listPngPath[int((i - 1) / 2)], width=Cm(7.2))
                            else:
                                paragraph.add_run(
                                    f"Figure {int(i / 2 + 1)}  {self.listTestInfoForReplace[2]} "
                                    f"{self.listTestInfoForReplace[0]}-{self.listTestInfoForReplace[1]} "
                                    f"Boxplot, {self.listCurrentLevel[int(i / 2 - 1)]}mA")
                        else:
                            paragraph._element.getparent().remove(paragraph._element)
                        continue

            # 识别插入点并插入表格
            if "Battery Quality Test / Alternative Battery Test for ESL Batteries" in paragraph.text:
                bInsertOverview = True
                intStepOut = 4
            elif "Version history" in paragraph.text and "Heading 2" == paragraph.style.name:
                bInsertVersionHistory = True
                intStepOut = 0
            elif "Test Information" in paragraph.text and "Heading 1" == paragraph.style.name:
                bInsertTestInformation = True
                intStepOut = 0
            elif "Test results" in paragraph.text and "Heading 1" == paragraph.style.name:
                bInsertStatisticalsResults = True
                intStepOut = 2

            if intStepOut:
                intStepOut = intStepOut - 1
            else:
                if bInsertOverview:
                    bInsertOverview = False
                    paragraph._p.addnext(tableOverview._tbl)
                elif bInsertVersionHistory:
                    bInsertVersionHistory = False
                    paragraph._p.addnext(tableVersionHistory._tbl)
                elif bInsertTestInformation:
                    bInsertTestInformation = False
                    paragraph._p.addnext(tableTestInformation._tbl)
                elif bInsertStatisticalsResults:
                    bInsertStatisticalsResults = False
                    paragraph._p.addnext(tableStatisticalsResults._tbl)

            # 删除温度符号
            if self.listTestInfo[7] == "Room Temperature" and "℃" in paragraph.text:
                paragraph.text = paragraph.text.replace("℃", "")

    # ── 主入口 ──

    def write(self, list_cpt=None, stats=None) -> None:
        """执行Word报告写入"""
        # 初始化Word文档
        wdReport = Document(self.strSampleReportWordPath)

        # 计算统计值
        if list_cpt is None:
            list_cpt = compute_list_cpt(
                self.listBatteryCharge, self.intBatteryNum,
                self.intCurrentLevelNum, self.intVoltageLevelNum)
        if stats is None:
            stats = compute_statistics(
                list_cpt, self.intCurrentLevelNum, self.intVoltageLevelNum)

        # 逐个写入各章节的表格
        tables = {
            'version_history': self._write_version_history(wdReport),
            'test_information': self._write_test_information(wdReport),
            'statistical_results': self._write_statistical_results(wdReport, stats),
        }

        # 准备Overview表内容（会修改 listTestInfoForReplace）
        overview_content = self._prepare_overview_content(wdReport, stats)
        tables['overview'] = self._write_overview_table(wdReport, overview_content, stats)

        # 文本替换、图片插入、表格插入
        self._replace_and_insert(wdReport, tables)

        # 清理并保存文档
        body = wdReport.element.body
        while len(body) > 0:
            last_child = body[-1]
            if last_child.tag.endswith('}p'):
                from docx.text.paragraph import Paragraph
                p = Paragraph(last_child, wdReport.element.body)
                if not p.text.strip():
                    body.remove(last_child)
                else:
                    break
            else:
                break
        wdReport.save(self.strReportWordPath)
        logging.info("数据分析完成，生成的docx报告路径: %s", self.strReportWordPath)
