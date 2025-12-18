import os
import csv
import sys
import json
import math
import datetime
import traceback
import configparser
import logging

# 导入软件版本信息
from battery_analysis import __version__
from battery_analysis.utils.config_utils import find_config_file

# 设置Matplotlib使用非交互式后端，避免线程安全问题
import matplotlib
matplotlib.use('Agg')  # 使用Agg后端，不会启动GUI

import numpy as np
import xlsxwriter as xwt
import matplotlib.pyplot as plt
from matplotlib.ticker import MultipleLocator
from openpyxl.utils import get_column_letter

from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_LINE_SPACING
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

from battery_analysis.utils.exception_type import BatteryAnalysisException
from battery_analysis.utils import numeric_utils
from battery_analysis.utils import excel_utils
from battery_analysis.utils import word_utils
from battery_analysis.utils import data_utils
from battery_analysis.utils import plot_utils
from battery_analysis.utils import csv_utils


class XlsxWordWriter:
    def __init__(self, strResultPath: str, listTestInfo: list, listBatteryInfo: list) -> None:
        # get config
        self.config = configparser.ConfigParser()
        
        # 使用通用配置文件查找函数
        config_path = find_config_file()
        
        # 尝试读取配置文件
        if config_path and os.path.exists(config_path):
            self.config.read(config_path, encoding='utf-8')
            logging.info(f"找到并读取配置文件: {config_path}")
        else:
            logging.warning("找不到配置文件，使用默认配置")
            # 创建默认的PltConfig部分
            self.config.add_section("PltConfig")
            self.config.set("PltConfig", "Title", "\"默认标题\"")
            # 创建默认的TestInformation部分
            self.config.add_section("TestInformation")
        
        # 安全获取PltConfig/Title，添加默认值处理
        try:
            if self.config.has_section("PltConfig") and self.config.has_option("PltConfig", "Title"):
                strImageTitle = self.config.get("PltConfig", "Title")[1:-1]
            else:
                # 使用测试信息生成默认标题
                strImageTitle = f"{listTestInfo[4]} {listTestInfo[2]} {listTestInfo[3]}({listTestInfo[5]}), 默认标题"
        except Exception as e:
            logging.error(f"获取标题时出错: {e}")
            strImageTitle = "默认电池分析图标题"

        # init variables for all files
        self.listTestInfo = listTestInfo
        self.listBatteryInfo = listBatteryInfo
        try:
            # 优先使用从Excel提取的Test Date（listBatteryInfo[3]）
            if len(self.listBatteryInfo) > 3 and self.listBatteryInfo[3] and self.listBatteryInfo[3] != "00000000":
                test_date = self.listBatteryInfo[3]
                logging.info(f"使用从Excel提取的Test Date: {test_date}")
                # 处理YYYYMMDD格式（8位数字）
                if len(test_date) == 8 and test_date.isdigit():
                    sy = test_date[:4]
                    sm = test_date[4:6]
                    sd = test_date[6:8]
                    td = f"{sy}{sm}{sd}"
                # 处理YYYY-MM-DD格式
                elif "-" in test_date:
                    [sy, sm, sd] = test_date.split(" ")[0].split("-")
                    td = f"{sy}{sm}{sd}"
                # 处理YYYY/MM/DD格式
                elif "/" in test_date:
                    [sy, sm, sd] = test_date.split(" ")[0].split("/")
                    td = f"{sy}{sm}{sd}"
                else:
                    raise ValueError(f"不支持的日期格式: {test_date}")
            else:
                # 使用从BatteryAnalysis类获取的test_date（在索引3位置）
                if len(self.listBatteryInfo) > 3 and self.listBatteryInfo[3] and self.listBatteryInfo[3] != "00000000":
                    # test_date已经是YYYYMMDD格式的字符串
                    td = self.listBatteryInfo[3]
                    logging.info(f"使用从Excel提取的Test Date: {td}")
                else:
                    raise ValueError("无法从BatteryInfo列表中提取有效日期信息")
            # 验证日期有效性
            if not (len(td) == 8 and td.isdigit() and td != "00000000"):
                raise ValueError(f"无效的日期格式: {td}")
            logging.info(f"成功解析日期: {td}")
        except (ValueError, IndexError) as e:
            logging.error(f"日期解析失败: {e}")
            # 尝试从文件名中提取日期
            if len(listTestInfo) > 0 and hasattr(listTestInfo[0], 'split'):
                import re
                filename = str(listTestInfo[0])
                # 匹配文件名中所有连续的数字组
                digit_groups = re.findall(r'(\d+)', filename)
                if digit_groups:
                    # 取最后一组连续数字
                    last_digit_group = digit_groups[-1]
                    # 提取前8位作为日期（如果长度足够）
                    if len(last_digit_group) >= 8:
                        td = last_digit_group[:8]
                        logging.info(f"从文件名最后一组连续数字提取前8位作为日期: {td}")
                    else:
                        # 如果最后一组数字不足8位，尝试匹配任意8位数字
                        match = re.search(r'(\d{8})', filename)
                        if match:
                            td = match.group(1)
                            logging.info(f"从文件名提取任意8位日期: {td}")
                        else:
                            td = "00000000"
                            logging.warning(f"无法从文件名提取日期，使用默认值: {td}")
                else:
                    td = "00000000"
                    logging.warning(f"文件名中没有数字，无法提取日期，使用默认值: {td}")
            else:
                td = "00000000"
                logging.warning(f"无法从文件名提取日期，使用默认值: {td}")
        # 使用os.path.join确保路径分隔符一致性
        self.strResultPath = os.path.join(strResultPath, f"{td}_V{listTestInfo[16]}")

        self.listCurrentLevel = listTestInfo[14]
        self.listVoltageLevel = listTestInfo[15]
        self.intCurrentLevelNum = len(self.listCurrentLevel)
        self.intVoltageLevelNum = len(self.listVoltageLevel)
        self.strFileCurrentType = ""
        for c in range(self.intCurrentLevelNum):
            self.strFileCurrentType = self.strFileCurrentType + f"{self.listCurrentLevel[c]}-"
        self.strFileCurrentType = self.strFileCurrentType[:-1]

        self.listBatteryCharge = self.listBatteryInfo[0]
        self.listBatteryName = self.listBatteryInfo[1]
        self.intBatteryNum = len(self.listBatteryName)

        # init variables for plt
        self.listBoxplotTitle = []
        self.listPngPath = []
        self.listSvgPath = []
        if len(strImageTitle) <= 70:
            strBoxplotTitle = strImageTitle
        else:
            strSplit = "A, "
            strBoxplotTitle = strImageTitle.split(strSplit)[0] + strSplit + "\n" + strImageTitle.split(strSplit)[1] + strSplit + strImageTitle.split(strSplit)[2]
        for b in range(self.intCurrentLevelNum):
            self.listBoxplotTitle.append(f"Useable Capacity over Cutoff Voltage, {self.listCurrentLevel[b]}mA Load\n{strBoxplotTitle}")
            self.listPngPath.append(f"{self.strResultPath}/Image_UseableCapacityOverCutoffVoltage{self.listCurrentLevel[b]}mALoad.png")
            self.listSvgPath.append(f"{self.strResultPath}/Image_UseableCapacityOverCutoffVoltage{self.listCurrentLevel[b]}mALoad.svg")
        self.strUnfilteredPngPath = f"{self.strResultPath}/Image_UnfilteredLoadVoltageOverCharge.png"
        self.strUnfilteredSvgPath = f"{self.strResultPath}/Image_UnfilteredLoadVoltageOverCharge.svg"
        self.strFilteredPngPath = f"{self.strResultPath}/Image_FilteredLoadVoltageOverCharge.png"
        self.strFilteredSvgPath = f"{self.strResultPath}/Image_FilteredLoadVoltageOverCharge.svg"
        self.strPltName = f"Load Voltage over Charge\n{strImageTitle}"
        self.strInfoImageCsvPath = f"{self.strResultPath}/Info_Image.csv"
        self.listPltColorType = ['#DF7040', '#0675BE', '#EDB120', '#7E2F8E', '#32CD32', '#FF4500', '#000000', '#000000']
        self.listColorName = ["red = ", "blue = ", "yellow = ", "violet = ", "green = ", "orange = ", "black1 = ", "black2 = "]
        strStrF = ""
        for c in range(self.intCurrentLevelNum):
            strStrF += f"{self.listColorName[c]}{self.listCurrentLevel[c]}mA, "
        strStrF = strStrF[:-2]

        # init variables for excel
        self.strResultXlsxPath = f"{self.strResultPath}/{self.listTestInfo[4]}_{self.listTestInfo[2]}_{self.listTestInfo[3]}_{self.strFileCurrentType}_{self.listTestInfo[7]}.xlsx"
        self.strSampleXlsxPath = f"{self.strResultPath}/Sample_{self.listTestInfo[4]}_{self.listTestInfo[2]}_{self.listTestInfo[3]}_{self.strFileCurrentType}_{self.listTestInfo[7]}.xlsx"

        # init variables for word
        if self.intCurrentLevelNum <= 4:
            self.strSampleReportWordPath = f"{self.strResultPath}/../../0_doc/Battery Measurement Report of TypeC TypeA_TypeD.docx"
        else:
            self.strSampleReportWordPath = f"{self.strResultPath}/../../0_doc/Battery Measurement Report of TypeC TypeA_TypeD_MP.docx"
            
        # strTimeStamp = datetime.datetime.now().strftime("%Y%m%d")

        # 使用pathlib.Path来规范化路径，避免出现../符号
        from pathlib import Path
        report_name = f"{self.listTestInfo[4]}_{self.listTestInfo[2]}_DC{self.listTestInfo[5]}_TD{td}_V{self.listTestInfo[16]}.docx"
        result_dir = Path(self.strResultPath).parent
        self.strReportWordPath = str(result_dir / report_name)
        self.listTextToReplace = ["TypeA", "TypeB", "TypeC", "TypeD", "TypeE", "TypeF", "TypeG", "StrA", "StrB", "StrC", "StrD", "StrF"]
        self.listImageToReplace = ["<<Image_FilteredLoadVoltageOverCharge>>"]
        for i in range(10):     # max 10 images to replace
            self.listImageToReplace.append(f"<<Image_UseableCapacityOverCutoffVoltage{i}>>")
            self.listImageToReplace.append(f"<<Title_UseableCapacityOverCutoffVoltage{i}>>")
        # 安全获取电池类型基础规格，添加默认值
        try:
            if self.config.has_section("BatteryConfig") and self.config.has_option("BatteryConfig", "SpecificationTypeBase"):
                listBatteryTypeBase = self.config.get("BatteryConfig", "SpecificationTypeBase").split(",")
            else:
                # 设置默认值
                listBatteryTypeBase = ["CoinCell", "ButtonCell", "Cylindrical", "Prismatic", "PouchCell"]
                logging.warning("使用默认电池类型基础规格")
                
            strBatteryType = ""
            for b in range(len(listBatteryTypeBase)):
                if listBatteryTypeBase[b].strip() in self.listTestInfo[2]:
                    strBatteryType = listBatteryTypeBase[b]
                    break
            
            # 如果没有匹配到，使用列表中的第一个或直接使用测试信息
            if strBatteryType == "":
                if listBatteryTypeBase:
                    strBatteryType = listBatteryTypeBase[0]
                    logging.warning(f"未找到精确匹配的电池类型，使用默认值: {strBatteryType}")
                else:
                    strBatteryType = self.listTestInfo[2]  # 直接使用测试信息中的类型
                    logging.warning(f"电池类型列表为空，直接使用测试信息: {strBatteryType}")
        except Exception as e:
            logging.error(f"获取电池类型时出错: {e}")
            strBatteryType = self.listTestInfo[2]  # 出错时使用测试信息中的类型
        self.listTestInfoForReplace = [self.listTestInfo[2], self.listTestInfo[3], self.listTestInfo[4],
                                       self.listTestInfo[5], self.listTestInfo[7], self.listTestInfo[11], strBatteryType,
                                       None, None, None, None, strStrF]

        # init variables for csv
        self.strResultCsvPath = f"{self.strResultPath}/{self.listTestInfo[4]}_{self.listTestInfo[2]}_{self.listTestInfo[3]}_{self.strFileCurrentType}_{self.listTestInfo[7]}.csv"

        # execute
        self.UXWW_XlsxWordCsvWrite()

    def UXWW_XlsxWordCsvWrite(self) -> None:
        # init excel writer
        wbResult = xwt.Workbook(self.strResultXlsxPath)
        wsOverview = wbResult.add_worksheet("overview")
        wsOverviewStatisticalStartLine = 13
        wsResult = wbResult.add_worksheet("result")
        wbSample = xwt.Workbook(self.strSampleXlsxPath)
        wsWord = wbSample.add_worksheet("word")
        wsExcel = wbSample.add_worksheet("excel")
        # init cxcel font format
        wsResultData = wbResult.add_format({
            'font_name': 'Microsoft YaHei',
            'font_size': 9,
            'font_color': 'black',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter'
        })
        wsResultData_italic = wbResult.add_format({
            'font_name': 'Microsoft YaHei',
            'font_size': 9,
            'font_color': 'black',
            'italic': True,
            'bold': False,
            'align': 'center',
            'valign': 'vcenter'
        })
        wsOverviewStatistics = wbResult.add_format({
            'font_name': 'Arial Narrow',
            'font_size': 9,
            'font_color': 'black',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': '#BFBFBF'
        })
        wsOverviewStatistics_bgdarkgray = wbResult.add_format({
            'font_name': 'Arial Narrow',
            'font_size': 12,
            'font_color': 'black',
            'bg_color': '#BFBFBF',
            'bold': True,
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#BFBFBF'
        })
        wsOverviewStatistics_bglightgray = wbResult.add_format({
            'font_name': 'Arial Narrow',
            'font_size': 12,
            'font_color': 'black',
            'bg_color': '#F2F2F2',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#BFBFBF'
        })
        wsOverviewStatistics_bglightgray_blod = wbResult.add_format({
            'font_name': 'Arial Narrow',
            'font_size': 12,
            'font_color': 'black',
            'bg_color': '#F2F2F2',
            'bold': True,
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#BFBFBF'
        })
        wsExcelLine = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': True,
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'border_color': 'black'
        })
        wsExcelData = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': 'black'
        })
        wsExcelData_bold = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': True,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': 'black'
        })
        wsExcelData_percentage = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'num_format': '0%',
            'border': 1,
            'border_color': 'black'
        })
        wsExcelData_percentage_bold = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': True,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'num_format': '0%',
            'border': 1,
            'border_color': 'black',
        })
        wsExcelData_bgyellow = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': 'black',
            'bg_color': '#FFFF00'
        })
        wsWordLine = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'border_color': 'black'
        })
        wsWordData = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': 'black'
        })
        wsWordData_bold = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': True,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': 'black'
        })
        wsWordData_percentage = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': 'black',
            'num_format': '0%'
        })
        wsWordData_percentage_bold = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': True,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': 'black',
            'num_format': '0%'
        })
        wsWordData_bgyellow = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': 'black',
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': 'black',
            'bg_color': '#FFFF00'
        })
        wbSampleHyperlink = wbSample.add_format({
            'font_name': 'Arial',
            'font_size': 10,
            'font_color': '#0000FF',
            'underline': True,
            'bold': False,
            'align': 'center',
            'valign': 'vcenter',
            'text_wrap': True,
            'border': 1,
            'border_color': 'black'
        })

        # Excel相关函数已移至excel_utils模块
        # 数值计算函数已移至numeric_utils模块

        # init word writer
        wdReport = Document(self.strSampleReportWordPath)

        # Word相关函数已移至word_utils模块

        # init csv writer
        f = open(self.strResultCsvPath, mode='w', newline='', encoding='utf-8')
        csvwriterResultCsvFile = csv.writer(f)
        
        # CSV写入缓冲区，减少I/O操作
        csv_buffer = []
        csv_buffer_size = 0
        max_csv_buffer_size = 100  # 每次写入100行

        # wbResult and csv write fixed part
        if wsOverview is None:
            raise BatteryAnalysisException(f"{self.strResultXlsxPath} sheet[overview] creation failed")
        wsOverview.write(0, 0, f"#BEGIN HEADER")
        wsOverview.write(1, 0, f"#PULSE DISCHARGE")
        wsOverview.write(2, 0, f"#BATTERY CHARACTERISTICS")
        wsOverview.write(3, 0, f"#Start Time: {self.listBatteryInfo[2][0]}")
        wsOverview.write(4, 0, f"#Start Time: {self.listBatteryInfo[2][1]}")
        wsOverview.write(5, 0, f"#Battery Type: {self.listTestInfo[2]} {self.listTestInfo[3]}")
        wsOverview.write(6, 0, f"#Battery Manufacturer: {self.listTestInfo[4]}")
        wsOverview.write(7, 0, f"#Battery Date Code: {self.listTestInfo[5]}")
        wsOverview.write(8, 0, f"#Temperature: {self.listTestInfo[6]}")
        wsOverview.write(9, 0, f"#Test Profile: {self.listTestInfo[13]}")
        wsOverview.write(10, 0, f"#Version: V1.0")
        wsOverview.write(11, 0, f"#END HEADER")
        csv_buffer_size = csv_utils.csv_write(f"#BEGIN HEADER", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#PULSE DISCHARGE", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#BATTERY CHARACTERISTICS", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#Start Time: {self.listBatteryInfo[2][0]}", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#End Time: {self.listBatteryInfo[2][1]}", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#Battery Type: {self.listTestInfo[2]} {self.listTestInfo[3]}", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#Battery Manufacturer: {self.listTestInfo[4]}", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#Battery Date Code: {self.listTestInfo[5]}", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#Temperature: {self.listTestInfo[6]}", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#Test Profile: {self.listTestInfo[13]}", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#Version: V1.0", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        csv_buffer_size = csv_utils.csv_write(f"#END HEADER", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)

        if wsResult is None:
            raise BatteryAnalysisException(f"{self.strResultXlsxPath} sheet[result] creation failed")
        excel_utils.ws_set_col(wsResult, 0, self.intCurrentLevelNum * (2 + self.intVoltageLevelNum) + 1, 10)
        excel_utils.ws_set_col(wsResult, 0, 1, 20)
        wsResult.write(2, 0, "Battery", wsResultData)
        wsResult.write(4 + self.intBatteryNum, 0, "Mean(\u03BC)", wsResultData_italic)
        wsResult.write(5 + self.intBatteryNum, 0, "Median", wsResultData_italic)
        wsResult.write(6 + self.intBatteryNum, 0, "Std. Var.(\u03C3)", wsResultData_italic)
        wsResult.write(7 + self.intBatteryNum, 0, "\u03BC-3\u03C3", wsResultData_italic)
        wsResult.write(8 + self.intBatteryNum, 0, "\u03BC-2\u03C3", wsResultData_italic)
        wsResult.write(9 + self.intBatteryNum, 0, "\u03BC+2\u03C3", wsResultData_italic)
        wsResult.write(10 + self.intBatteryNum, 0, "\u03BC+3\u03C3", wsResultData_italic)
        wsResult.write(11 + self.intBatteryNum, 0, "Minimum", wsResultData_italic)
        wsResult.write(12 + self.intBatteryNum, 0, "Maximum", wsResultData_italic)
        for c in range(self.intCurrentLevelNum):
            wsResult.write(1, 1 + c * (2 + self.intVoltageLevelNum), f"{self.listCurrentLevel[c]}mA", wsResultData)
            wsResult.merge_range(1, 2 + c * (2 + self.intVoltageLevelNum), 1, (c + 1) * (2 + self.intVoltageLevelNum) - 1, "Voltage", wsResultData)
            for v in range(self.intVoltageLevelNum):
                wsResult.write(2, 2 + c * (2 + self.intVoltageLevelNum) + v, f"{self.listVoltageLevel[v]}V", wsResultData)

        csv_buffer_size = csv_utils.csv_write("", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        listCsvLine = [""]
        for c in range(self.intCurrentLevelNum):
            listCsvLine.append(f"{self.listCurrentLevel[c]}mA")
            listCsvLine.append("Voltage")
            for v in range(self.intVoltageLevelNum):
                listCsvLine.append("")
        csv_buffer_size = csv_utils.csv_write(listCsvLine, csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        listCsvLine = []
        for c in range(self.intCurrentLevelNum):
            listCsvLine.append("")
            listCsvLine.append("")
            for v in range(self.intVoltageLevelNum):
                listCsvLine.append(f"{self.listVoltageLevel[v]}V")
        listCsvLine[0] = "Battery"
        csv_buffer_size = csv_utils.csv_write(listCsvLine, csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)

        # wdReport write table Version History
        tableVersionHistory = wdReport.add_table(9, 4, style='Grid Table 4 Accent 3')
        for r in range(9):
            tableVersionHistory.rows[r].height = Cm(0.6)
            for c in range(4):
                if c < 4:
                    tableVersionHistory.cell(0, c).width = Cm(2.8)
                cell = tableVersionHistory.cell(r, c)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        tableVersionHistory.cell(0, 0).paragraphs[0].add_run("Date").font.size = Pt(10)
        tableVersionHistory.cell(0, 1).paragraphs[0].add_run("Version").font.size = Pt(10)
        tableVersionHistory.cell(0, 2).paragraphs[0].add_run("Editor").font.size = Pt(10)
        tableVersionHistory.cell(0, 3).paragraphs[0].add_run("Changes").font.size = Pt(10)
        tableVersionHistory.cell(1, 1).paragraphs[0].add_run("1.0").font.size = Pt(10)
        tableVersionHistory.cell(1, 3).paragraphs[0].add_run("Initial version").font.size = Pt(10)
        text = tableVersionHistory.cell(1, 0).paragraphs[0].add_run(datetime.datetime.now().strftime("%Y.%m.%d"))
        text.font.size = Pt(10)
        text.font.bold = False
        strReportedBy = self.listTestInfo[18] if len(self.listTestInfo) > 18 else ""
        tableVersionHistory.cell(1, 2).paragraphs[0].add_run(f"{strReportedBy}").font.size = Pt(10)

        # wdReport write table Test Information
        tableTestInformation = wdReport.add_table(5, 2, style='Table Grid')
        tableTestInformation.cell(0, 1).width = Cm(10)
        for r in range(5):
            for c in range(2):
                cell = tableTestInformation.cell(r, c)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
                cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        tableTestInformation.cell(0, 0).paragraphs[0].add_run("Test Equipment").font.size = Pt(10)
        tableTestInformation.cell(1, 0).paragraphs[0].add_run("Software Versions").font.size = Pt(10)
        tableTestInformation.cell(2, 0).paragraphs[0].add_run("Middle Machines").font.size = Pt(10)
        tableTestInformation.cell(3, 0).paragraphs[0].add_run("Test Units").font.size = Pt(10)
        tableTestInformation.cell(4, 0).paragraphs[0].add_run("Data Processing Platforms").font.size = Pt(10)

        strItemTestEquipment = word_utils.get_item(self.config, "TestInformation", "TestEquipment")
        tableTestInformation.cell(0, 1).paragraphs[0].add_run(f"{strItemTestEquipment}").font.size = Pt(10)
        strItemSoftwareVersionsBTSServerVersion = word_utils.get_item(self.config, "TestInformation", "SoftwareVersions.BTSServerVersion")
        strItemSoftwareVersionsBTSClientVersion = word_utils.get_item(self.config, "TestInformation", "SoftwareVersions.BTSClientVersion")
        strItemSoftwareVersionsTSDAVersion = word_utils.get_item(self.config, "TestInformation", "SoftwareVersions.BTSDAVersion")
        tableTestInformation.cell(1, 1).paragraphs[0].add_run(f"BTS Server Version: {strItemSoftwareVersionsBTSServerVersion}\n"
                                                              f"BTS Client Version: {strItemSoftwareVersionsBTSClientVersion}\n"
                                                              f"BTSDA (Data Analysis) Version: {strItemSoftwareVersionsTSDAVersion}").font.size = Pt(10)
        strItemMiddleMachinesModel = word_utils.get_item(self.config, "TestInformation", "MiddleMachines.Model", 6)
        strItemMiddleMachinesHardwareVersion = word_utils.get_item(self.config, "TestInformation", "MiddleMachines.HardwareVersion", 14)
        strItemMiddleMachinesSerialNumber = word_utils.get_item(self.config, "TestInformation", "MiddleMachines.SerialNumber", 12)
        strItemMiddleMachinesFirmwareVersion = word_utils.get_item(self.config, "TestInformation", "MiddleMachines.FirmwareVersion", 14)
        strItemMiddleMachinesDeviceType = word_utils.get_item(self.config, "TestInformation", "MiddleMachines.DeviceType", 10)
        tableTestInformation.cell(2, 1).paragraphs[0].add_run(f"Model: {strItemMiddleMachinesModel}\n"
                                                              f"Hardware Version: {strItemMiddleMachinesHardwareVersion}\n"
                                                              f"Serial Number: {strItemMiddleMachinesSerialNumber}\n"
                                                              f"Firmware Version: {strItemMiddleMachinesFirmwareVersion}\n"
                                                              f"Device Type: {strItemMiddleMachinesDeviceType}").font.size = Pt(10)
        strItemTestUnitsModel = word_utils.get_item(self.config, "TestInformation", "TestUnits.Model", 6)
        strItemTestUnitsHardwareVersion = word_utils.get_item(self.config, "TestInformation", "TestUnits.HardwareVersion", 14)
        strItemTestUnitsFirmwareVersion = word_utils.get_item(self.config, "TestInformation", "TestUnits.FirmwareVersion", 14)
        tableTestInformation.cell(3, 1).paragraphs[0].add_run(f"Model: {strItemTestUnitsModel}\n"
                                                              f"Hardware Version: {strItemTestUnitsHardwareVersion}\n"
                                                              f"Firmware Version: {strItemTestUnitsFirmwareVersion}").font.size = Pt(10)
        # Data Processing Platforms 跟随软件版本，不再从配置文件读取
        strItemDataProcessingPlatforms = f"Battery Analyzer-v{__version__}"
        tableTestInformation.cell(4, 1).paragraphs[0].add_run(f"{strItemDataProcessingPlatforms}").font.size = Pt(10)

        # wbResult and csv write analytical battery statistic
        for b in range(self.intBatteryNum):
            excel_utils.ws_result_write_data(3 + b, 0, self.listBatteryName[b], wsResultData, wsResult)
            listCsvLine = []
            i = 0
            for c in range(self.intCurrentLevelNum):
                listCsvLine.append("")
                listCsvLine.append("")
                for v in range(self.intVoltageLevelNum):
                    listCsvLine.append(self.listBatteryCharge[b][i])
                    excel_utils.ws_result_write_data(3 + b, 2 + c * (2 + self.intVoltageLevelNum) + v, self.listBatteryCharge[b][i], wsResultData, wsResult)
                    i += 1
            listCsvLine[0] = f"{self.listBatteryName[b]}"
            csv_buffer_size = csv_utils.csv_write(listCsvLine, csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)

        # init and fill listCpt for calculate
        listCpt = []
        for c in range(self.intCurrentLevelNum):
            listCpt.append([])
            for _ in range(self.intVoltageLevelNum):
                listCpt[c].append([])
        for b in range(self.intBatteryNum):
            i = 0
            for c in range(self.intCurrentLevelNum):
                for v in range(self.intVoltageLevelNum):
                    if self.listBatteryCharge[b][i] != 0:
                        listCpt[c][v].append(self.listBatteryCharge[b][i])
                    i += 1

        # draw boxplots
        self.UXWW_Draw(listCpt, int(self.listTestInfo[8]))

        # init list for calculate
        listMean = []
        listStd = []
        listMax = []
        listMin = []
        listMed = []
        listMM3S = []
        listMM2S = []
        listMP2S = []
        listMP3S = []
        for c in range(self.intCurrentLevelNum):
            listMean.append([])
            listMed.append([])
            listStd.append([])
            listMM3S.append([])
            listMM2S.append([])
            listMP2S.append([])
            listMP3S.append([])
            listMin.append([])
            listMax.append([])

        # calculate
        for c in range(self.intCurrentLevelNum):
            for v in range(self.intVoltageLevelNum):
                listMean[c].append(numeric_utils.np_mean(listCpt[c][v]))
                listMed[c].append(numeric_utils.np_med(listCpt[c][v]))
                listStd[c].append(numeric_utils.np_std(listCpt[c][v]))
                listMM3S[c].append(listMean[c][v] - 3 * listStd[c][v])
                listMM2S[c].append(listMean[c][v] - 2 * listStd[c][v])
                listMP2S[c].append(listMean[c][v] + 2 * listStd[c][v])
                listMP3S[c].append(listMean[c][v] + 3 * listStd[c][v])
                listMin[c].append(numeric_utils.np_min(listCpt[c][v]))
                listMax[c].append(numeric_utils.np_max(listCpt[c][v]))

        # wbResult and csv write calculated statistic
        listCsvName = ["Mean(\u03BC)", "Median", "Std. Var.(\u03C3)", "\u03BC-3\u03C3", "\u03BC-2\u03C3", "\u03BC+2\u03C3", "\u03BC+3\u03C3", "Minimum", "Maximum"]
        listCsvList = [listMean, listMed, listStd, listMM3S, listMM2S, listMP2S, listMP3S, listMin, listMax]
        csv_buffer_size = csv_utils.csv_write("", csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)
        for n in range(len(listCsvName)):
            listCsvLine = []
            for c in range(self.intCurrentLevelNum):
                listCsvLine.append("")
                listCsvLine.append("")
                for v in range(self.intVoltageLevelNum):
                    excel_utils.ws_result_write_data(4 + self.intBatteryNum, 2 + c * (2 + self.intVoltageLevelNum) + v, round(listMean[c][v], 5), wsResultData, wsResult)
                    excel_utils.ws_result_write_data(5 + self.intBatteryNum, 2 + c * (2 + self.intVoltageLevelNum) + v, round(listMed[c][v], 5), wsResultData, wsResult)
                    excel_utils.ws_result_write_data(6 + self.intBatteryNum, 2 + c * (2 + self.intVoltageLevelNum) + v, round(listStd[c][v], 5), wsResultData, wsResult)
                    excel_utils.ws_result_write_data(7 + self.intBatteryNum, 2 + c * (2 + self.intVoltageLevelNum) + v, round(listMM3S[c][v], 5), wsResultData, wsResult)
                    excel_utils.ws_result_write_data(8 + self.intBatteryNum, 2 + c * (2 + self.intVoltageLevelNum) + v, round(listMM2S[c][v], 5), wsResultData, wsResult)
                    excel_utils.ws_result_write_data(9 + self.intBatteryNum, 2 + c * (2 + self.intVoltageLevelNum) + v, round(listMP2S[c][v], 5), wsResultData, wsResult)
                    excel_utils.ws_result_write_data(10 + self.intBatteryNum, 2 + c * (2 + self.intVoltageLevelNum) + v, round(listMP3S[c][v], 5), wsResultData, wsResult)
                    excel_utils.ws_result_write_data(11 + self.intBatteryNum, 2 + c * (2 + self.intVoltageLevelNum) + v, round(listMin[c][v], 5), wsResultData, wsResult)
                    excel_utils.ws_result_write_data(12 + self.intBatteryNum, 2 + c * (2 + self.intVoltageLevelNum) + v, round(listMax[c][v], 5), wsResultData, wsResult)
                    listCsvLine.append(round(listCsvList[n][c][v], 5))
            listCsvLine[0] = f"{listCsvName[n]}"
            csv_buffer_size = csv_utils.csv_write(listCsvLine, csvwriterResultCsvFile, csv_buffer, csv_buffer_size, max_csv_buffer_size)

        for c in range(self.intCurrentLevelNum):
            wsResult.insert_image(14 + self.intBatteryNum, 1 + c * (2 + self.intVoltageLevelNum), self.listPngPath[c], {'x_scale': ((2 + self.intVoltageLevelNum) * 2 - 1) / 16, 'y_scale': ((2 + self.intVoltageLevelNum) * 2 - 1) / 16})

        tableStatisticalsResults = wdReport.add_table(self.intCurrentLevelNum + 1, self.intVoltageLevelNum + 1, style='Table Grid')
        for c in range(self.intCurrentLevelNum + 1):
            for v in range(self.intVoltageLevelNum + 1):
                cell = tableStatisticalsResults.cell(c, v)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                if c == 0 and v == 0:
                    tableStatisticalsResults.rows[c].height = Cm(0.7)
                    tableStatisticalsResults.cell(c, v).width = Cm(3.55)
                    text = cell.paragraphs[0].add_run("Statisticals\nResults")
                    text.font.size = Pt(12)
                    text.bold = True
                    word_utils.table_set_bg_color(cell, '#BFBFBF')
                    wsOverview.set_row(wsOverviewStatisticalStartLine, 20)
                    wsOverview.set_column(v, v, 18)
                    wsOverview.write(wsOverviewStatisticalStartLine, v, "Statisticals Results", wsOverviewStatistics_bgdarkgray)
                elif c == 0 and v > 0:
                    tableStatisticalsResults.cell(c, v).width = Cm(3.55)
                    text1 = cell.paragraphs[0].add_run(f"Cut-off Voltage\n")
                    text1.font.size = Pt(12)
                    text2 = cell.paragraphs[0].add_run(f"{self.listVoltageLevel[v - 1]}V")
                    text2.font.bold = True
                    text2.font.size = Pt(12)
                    word_utils.table_set_bg_color(cell, '#F2F2F2')
                    wsOverview.set_column(v, v, 18)
                    wsOverview.write_rich_string(wsOverviewStatisticalStartLine, v,
                                                 wsOverviewStatistics_bglightgray, "Cut-off Voltage ",
                                                 wsOverviewStatistics_bglightgray_blod, f"{self.listVoltageLevel[v - 1]}V",
                                                 wsOverviewStatistics_bglightgray)
                elif c > 0 and v == 0:
                    tableStatisticalsResults.rows[c].height = Cm(2.35)
                    text1 = cell.paragraphs[0].add_run(f"Pulse Current\n")
                    text1.font.size = Pt(12)
                    text2 = cell.paragraphs[0].add_run(f"{self.listCurrentLevel[c - 1]}mA")
                    text2.font.bold = True
                    text2.font.size = Pt(12)
                    word_utils.table_set_bg_color(cell, '#F2F2F2')
                    wsOverview.set_row(wsOverviewStatisticalStartLine + c, 120)
                    wsOverview.write_rich_string(wsOverviewStatisticalStartLine + c, 0,
                                                 wsOverviewStatistics_bglightgray, "Pulse Current ",
                                                 wsOverviewStatistics_bglightgray_blod, f"{self.listCurrentLevel[c - 1]}mA",
                                                 wsOverviewStatistics_bglightgray)
                else:
                    text = cell.paragraphs[0].add_run(f"\u03BC: {round(listMean[c - 1][v - 1])}mAh\n"
                                                      f"Median: {round(listMed[c - 1][v - 1])}mAh\n"
                                                      f"\u03C3: {round(listStd[c - 1][v - 1])}mAh\n"
                                                      f"\u03BC - 3\u03C3: {round(listMM3S[c - 1][v - 1])}mAh\n"
                                                      f"\u03BC - 2\u03C3: {round(listMM2S[c - 1][v - 1])}mAh\n"
                                                      f"\u03BC + 2\u03C3: {round(listMP2S[c - 1][v - 1])}mAh\n"
                                                      f"\u03BC + 3\u03C3: {round(listMP3S[c - 1][v - 1])}mAh\n"
                                                      f"Minimum: {round(listMin[c - 1][v - 1])}mAh\n"
                                                      f"Maximum: {round(listMax[c - 1][v - 1])}mAh")
                    text.font.size = Pt(7)
                    cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.EXACTLY
                    cell.paragraphs[0].paragraph_format.line_spacing = Pt(10)
                    wsOverview.write(wsOverviewStatisticalStartLine + c, v, f"\u03BC: {round(listMean[c - 1][v - 1])}mAh\n"
                                                                            f"Median: {round(listMed[c - 1][v - 1])}mAh\n"
                                                                            f"\u03C3: {round(listStd[c - 1][v - 1])}mAh\n"
                                                                            f"\u03BC - 3\u03C3: {round(listMM3S[c - 1][v - 1])}mAh\n"
                                                                            f"\u03BC - 2\u03C3: {round(listMM2S[c - 1][v - 1])}mAh\n"
                                                                            f"\u03BC + 2\u03C3: {round(listMP2S[c - 1][v - 1])}mAh\n"
                                                                            f"\u03BC + 3\u03C3: {round(listMP3S[c - 1][v - 1])}mAh\n"
                                                                            f"Minimum: {round(listMin[c - 1][v - 1])}mAh\n"
                                                                            f"Maximum: {round(listMax[c - 1][v - 1])}mAh",
                                     wsOverviewStatistics)

        # wbSample write data
        if self.listTestInfo[0] == "Coin Cell":
            intTestProfileStartLine = 3
        elif self.listTestInfo[0] == "Pouch Cell":
            intTestProfileStartLine = 4
        else:
            raise BatteryAnalysisException("[Test Info Error]: listTestInfo[0] is a unknown battery type")
        excel_utils.ws_set_col(wsExcel, 0, 3, 12)
        excel_utils.ws_set_col(wsExcel, 3, 1, 20)
        excel_utils.ws_set_col(wsExcel, intTestProfileStartLine, 1, 12)
        excel_utils.ws_set_col(wsExcel, intTestProfileStartLine + 1, 1, 15)
        excel_utils.ws_set_col(wsExcel, intTestProfileStartLine + 2, 1, 10)
        excel_utils.ws_set_col(wsExcel, intTestProfileStartLine + 3, 1, 18)
        excel_utils.ws_set_col(wsExcel, intTestProfileStartLine + 4, 1, 25)
        excel_utils.ws_set_col(wsExcel, intTestProfileStartLine + 5, 2, 30)
        excel_utils.ws_set_col(wsExcel, intTestProfileStartLine + 7, 2, 15)
        intActualMeasuredCapacityLength = self.intVoltageLevelNum * 2
        excel_utils.ws_set_col(wsExcel, intTestProfileStartLine + 9, intActualMeasuredCapacityLength, 6)
        intTestDateStartCol = intTestProfileStartLine + 9 + intActualMeasuredCapacityLength
        excel_utils.ws_set_col(wsExcel, intTestDateStartCol, 1, 10)
        excel_utils.ws_set_col(wsExcel, intTestDateStartCol + 1, 1, 12)
        excel_utils.ws_set_col(wsExcel, intTestDateStartCol + 2, 1, 18)
        excel_utils.ws_set_col(wsExcel, intTestDateStartCol + 3, 1, 8)
        excel_utils.ws_set_col(wsExcel, intTestDateStartCol + 4, 2, 40)

        intPosiMaxmA = 0
        intPosi2V25 = 0
        intPresentmA = 0
        for c in range(self.intCurrentLevelNum):
            intPresentmA = self.listCurrentLevel[c]
            if self.listCurrentLevel[c] > intPresentmA:
                intPosiMaxmA = c
                break
        for v in range(self.intVoltageLevelNum):
            if self.listVoltageLevel[v] == 2.25:
                intPosi2V25 = v
                break

        listStrItems = [
            "Battery Type",  # 0
            "Specification",  # 1
            "Manufacturer",  # 2
            "Construction Method",  # 3
            "Test Profile",  # 4
            "Tester location",  # 5
            "Tested By",  # 6
            "Batch/Date Code",  # 7
            "Accelerated Aging[Years]",  # 8
            "Datasheet Nominal Capacity[mAh]",  # 9
            "Calculation Nominal Capacity[mAh]",  # 10
            "Required Useable Capacity[mAh]",  # 11
            None,  # 12
            f"Actual Measured Capacity[mAh]\n(at {self.listCurrentLevel[intPosiMaxmA]}mA/2.25V)",  # 13
            "Test Date",  # 14
            "Samples Qty",  # 15
            "Temperature[\u2103]",  # 16
            "Result",  # 17
            "Test Results File",  # 18
            "Remarks"  # 19
        ]

        try:
            strRelProfilePath = os.path.relpath(self.listTestInfo[13], os.path.dirname(self.strSampleXlsxPath))
        except ValueError:
            strRelProfilePath = self.listTestInfo[13]

        if self.listTestInfo[5] == "":
            strBatchDateCode = "n.a."
        else:
            strBatchDateCode = self.listTestInfo[5]

        intPassRate = float(int(self.listTestInfo[17])/int(self.listTestInfo[9]))
        strRequiredUseableCapacityPercentage = f"{int(100*int(self.listTestInfo[17])/int(self.listTestInfo[9]))}%"
        self.listTestInfoForReplace[9] = strRequiredUseableCapacityPercentage

        [sy, sm, sd] = self.listBatteryInfo[2][0].split(" ")[0].split("-")
        [ey, em, ed] = self.listBatteryInfo[2][1].split(" ")[0].split("-")

        strRelResultPath = os.path.relpath(self.strResultXlsxPath, os.path.dirname(self.strReportWordPath))
        if listMM2S[intPosiMaxmA][intPosi2V25] / int(self.listTestInfo[9]) >= intPassRate:
            strResult = "Pass"
            self.listTestInfoForReplace[8] = "meets"
            self.listTestInfoForReplace[10] = "Pass"
            strRemarks = "OK"
        else:
            strResult = "Fail"
            self.listTestInfoForReplace[8] = "doesn't meet"
            self.listTestInfoForReplace[10] = "Fail"
            strRemarks = f"The expected usable Q(stat) should be more than {strRequiredUseableCapacityPercentage}, " \
                         f"while the actual measured minimum capacity to 2.25V is " \
                         f"{math.floor(100 * listMM2S[intPosiMaxmA][intPosi2V25] / int(self.listTestInfo[9]))}%."
        self.listTestInfoForReplace[7] = f"{math.floor(100 * listMM2S[intPosiMaxmA][intPosi2V25] / int(self.listTestInfo[9]))}%"

        listStrContent = [
            self.listTestInfo[0],  # 0
            f"{self.listTestInfo[2]}-{self.listTestInfo[3]}",  # 1
            self.listTestInfo[4],  # 2
            self.listTestInfo[1],  # 3
            strRelProfilePath,  # 4
            self.listTestInfo[11],  # 5
            self.listTestInfo[12],  # 6
            strBatchDateCode,  # 7
            self.listTestInfo[10],  # 8
            self.listTestInfo[8],  # 9
            self.listTestInfo[9],  # 10
            self.listTestInfo[17],  # 11
            strRequiredUseableCapacityPercentage,  # 12
            None,  # 13
            f"{sd}.{sm}.{sy} - {ed}.{em}.{ey}",  # 14
            self.listTestInfo[6],  # 15
            self.listTestInfo[7],  # 16
            strResult,  # 17
            strRelResultPath,  # 18
            strRemarks  # 19
        ]

        wsExcel.merge_range(0, 0, 1, 0, listStrItems[0], wsExcelLine)
        wsExcel.merge_range(0, 1, 1, 1, listStrItems[1], wsExcelLine)
        wsExcel.merge_range(0, 2, 1, 2, listStrItems[2], wsExcelLine)
        wsExcel.merge_range(0, 3, 1, 3, listStrItems[3], wsExcelLine)
        if intTestProfileStartLine == 4:
            wsExcel.merge_range(0, intTestProfileStartLine, 1, intTestProfileStartLine, "", wsExcelLine)
        wsExcel.write(0, intTestProfileStartLine, listStrItems[4], wsExcelLine)
        wsExcel.merge_range(0, intTestProfileStartLine + 1, 1, intTestProfileStartLine + 1, listStrItems[5], wsExcelLine)
        wsExcel.merge_range(0, intTestProfileStartLine + 2, 1, intTestProfileStartLine + 2, listStrItems[6], wsExcelLine)
        wsExcel.merge_range(0, intTestProfileStartLine + 3, 1, intTestProfileStartLine + 3, listStrItems[7], wsExcelLine)
        wsExcel.merge_range(0, intTestProfileStartLine + 4, 1, intTestProfileStartLine + 4, listStrItems[8], wsExcelLine)
        wsExcel.merge_range(0, intTestProfileStartLine + 5, 1, intTestProfileStartLine + 5, listStrItems[9], wsExcelLine)
        wsExcel.merge_range(0, intTestProfileStartLine + 6, 1, intTestProfileStartLine + 6, listStrItems[10], wsExcelLine)
        wsExcel.merge_range(0, intTestProfileStartLine + 7, 1, intTestProfileStartLine + 8, listStrItems[11], wsExcelLine)
        wsExcel.merge_range(0, intTestProfileStartLine + 9, 0, intTestDateStartCol - 1, listStrItems[13], wsExcelLine)
        for v in range(self.intVoltageLevelNum):
            wsExcel.merge_range(1, intTestProfileStartLine + 9 + v * 2, 1, intTestProfileStartLine + 9 + v * 2 + 1, f"{self.listVoltageLevel[v]}V", wsExcelLine)
        wsExcel.merge_range(0, intTestDateStartCol, 1, intTestDateStartCol, listStrItems[14], wsExcelLine)
        wsExcel.merge_range(0, intTestDateStartCol + 1, 1, intTestDateStartCol + 1, listStrItems[15], wsExcelLine)
        wsExcel.merge_range(0, intTestDateStartCol + 2, 1, intTestDateStartCol + 2, listStrItems[16], wsExcelLine)
        wsExcel.merge_range(0, intTestDateStartCol + 3, 1, intTestDateStartCol + 3, listStrItems[17], wsExcelLine)
        wsExcel.merge_range(0, intTestDateStartCol + 4, 1, intTestDateStartCol + 4, listStrItems[18], wsExcelLine)
        wsExcel.merge_range(0, intTestDateStartCol + 5, 1, intTestDateStartCol + 5, listStrItems[19], wsExcelLine)

        wsExcel.write(2, 0, listStrContent[0], wsExcelData)
        wsExcel.write(2, 1, listStrContent[1], wsExcelData)
        wsExcel.write(2, 2, listStrContent[2], wsExcelData)
        wsExcel.write(2, 3, listStrContent[3], wsExcelData)
        if len(listStrContent[4].split("\\")) == 1:
            wsExcel.write(2, intTestProfileStartLine, listStrContent[4], wsExcelData)
        else:
            # 将相对路径转换为file:// URL格式
            url_path = listStrContent[4]
            # 确保路径使用正斜杠
            url_path = url_path.replace('\\', '/')
            # 添加file://前缀
            file_url = f'file:///{url_path}'
            wsExcel.write_url(2, intTestProfileStartLine, file_url, wbSampleHyperlink, string=listStrContent[4].split("\\")[-1])
        wsExcel.write(2, intTestProfileStartLine + 1, listStrContent[5], wsExcelData)
        wsExcel.write(2, intTestProfileStartLine + 2, listStrContent[6], wsExcelData)
        wsExcel.write(2, intTestProfileStartLine + 3, listStrContent[7], wsExcelData)
        wsExcel.write(2, intTestProfileStartLine + 4, listStrContent[8], wsExcelData)
        wsExcel.write(2, intTestProfileStartLine + 5, listStrContent[9], wsExcelData)
        wsExcel.write(2, intTestProfileStartLine + 6, listStrContent[10], wsExcelData)
        wsExcel.write(2, intTestProfileStartLine + 7, listStrContent[11], wsExcelData)
        wsExcel.write(2, intTestProfileStartLine + 8, listStrContent[12], wsExcelData)

        for v in range(self.intVoltageLevelNum):
            if v == intPosi2V25:
                wsExcel.write(2, intTestProfileStartLine + 9 + v * 2, f"{round(listMM2S[intPosiMaxmA][v], 2)}", wsExcelData_bold)
                wsExcel.write_formula(f"{excel_utils.num2letter(intTestProfileStartLine + 9 + v * 2 + 1)}3",
                                      f"=TRUNC({excel_utils.num2letter(intTestProfileStartLine + 9 + v * 2)}3/{excel_utils.num2letter(intTestProfileStartLine + 6)}3, 2)",
                                      wsExcelData_percentage_bold)
            else:
                wsExcel.write(2, intTestProfileStartLine + 9 + v * 2, f"{round(listMM2S[intPosiMaxmA][v], 2)}", wsExcelData)
                wsExcel.write_formula(f"{excel_utils.num2letter(intTestProfileStartLine + 9 + v * 2 + 1)}3",
                                      f"=TRUNC({excel_utils.num2letter(intTestProfileStartLine + 9 + v * 2)}3/{excel_utils.num2letter(intTestProfileStartLine + 6)}3, 2)",
                                      wsExcelData_percentage)

        wsExcel.write(2, intTestDateStartCol, listStrContent[14], wsExcelData)
        wsExcel.write(2, intTestDateStartCol + 1, listStrContent[15], wsExcelData)
        wsExcel.write(2, intTestDateStartCol + 2, listStrContent[16], wsExcelData)
        wsExcel.write(2, intTestDateStartCol + 3, listStrContent[17], wsExcelData_bgyellow)
        # 将相对路径转换为file:// URL格式
        url_path = listStrContent[18]
        # 确保路径使用正斜杠
        url_path = url_path.replace('\\', '/')
        # 添加file://前缀
        file_url = f'file:///{url_path}'
        wsExcel.write_url(2, intTestDateStartCol + 4, file_url, wbSampleHyperlink, string=listStrContent[18].split("\\")[-1])
        wsExcel.write(2, intTestDateStartCol + 5, listStrContent[19], wsExcelData)

        excel_utils.ws_set_col(wsWord, 0, 1, 30)
        excel_utils.ws_set_col(wsWord, 1, intActualMeasuredCapacityLength, 3)
        wsWord.write(0, 0, listStrItems[0], wsWordLine)
        wsWord.write(1, 0, listStrItems[1], wsWordLine)
        wsWord.write(2, 0, listStrItems[2], wsWordLine)
        wsWord.write(3, 0, listStrItems[3], wsWordLine)
        wsWord.write(intTestProfileStartLine, 0, listStrItems[4], wsWordLine)
        wsWord.write(intTestProfileStartLine + 1, 0, listStrItems[5], wsWordLine)
        wsWord.write(intTestProfileStartLine + 2, 0, listStrItems[6], wsWordLine)
        wsWord.write(intTestProfileStartLine + 3, 0, listStrItems[7], wsWordLine)
        wsWord.write(intTestProfileStartLine + 4, 0, listStrItems[8], wsWordLine)
        wsWord.write(intTestProfileStartLine + 5, 0, listStrItems[9], wsWordLine)
        wsWord.write(intTestProfileStartLine + 6, 0, listStrItems[10], wsWordLine)
        wsWord.write(intTestProfileStartLine + 7, 0, listStrItems[11], wsWordLine)
        wsWord.merge_range(intTestProfileStartLine + 8, 0, intTestProfileStartLine + 10, 0, listStrItems[13], wsWordLine)
        intTestDateStartRow = intTestProfileStartLine + 11
        wsWord.write(intTestDateStartRow, 0, listStrItems[14], wsWordLine)
        wsWord.write(intTestDateStartRow + 1, 0, listStrItems[15], wsWordLine)
        wsWord.write(intTestDateStartRow + 2, 0, listStrItems[16], wsWordLine)
        wsWord.write(intTestDateStartRow + 3, 0, listStrItems[17], wsWordLine)
        wsWord.write(intTestDateStartRow + 4, 0, listStrItems[18], wsWordLine)
        wsWord.write(intTestDateStartRow + 5, 0, listStrItems[19], wsWordLine)

        wsWord.merge_range(0, 1, 0, intActualMeasuredCapacityLength, listStrContent[0], wsWordData)
        wsWord.merge_range(1, 1, 1, intActualMeasuredCapacityLength, listStrContent[1], wsWordData)
        wsWord.merge_range(2, 1, 2, intActualMeasuredCapacityLength, listStrContent[2], wsWordData)
        wsWord.merge_range(3, 1, 3, intActualMeasuredCapacityLength, listStrContent[3], wsWordData)
        if intTestProfileStartLine == 4:
            wsWord.merge_range(intTestProfileStartLine, 1, intTestProfileStartLine, intActualMeasuredCapacityLength, "", wsWordData)
        if len(listStrContent[4].split("\\")) == 1:
            wsWord.write(intTestProfileStartLine, 1, listStrContent[4], wsWordData)
        else:
            # 将相对路径转换为file:// URL格式
            url_path = listStrContent[4]
            # 确保路径使用正斜杠
            url_path = url_path.replace('\\', '/')
            # 添加file://前缀
            file_url = f'file:///{url_path}'
            wsWord.write_url(intTestProfileStartLine, 1, file_url, wbSampleHyperlink, string=listStrContent[4].split("\\")[-1])
        wsWord.merge_range(intTestProfileStartLine + 1, 1, intTestProfileStartLine + 1, intActualMeasuredCapacityLength, listStrContent[5], wsWordData)
        wsWord.merge_range(intTestProfileStartLine + 2, 1, intTestProfileStartLine + 2, intActualMeasuredCapacityLength, listStrContent[6], wsWordData)
        wsWord.merge_range(intTestProfileStartLine + 3, 1, intTestProfileStartLine + 3, intActualMeasuredCapacityLength, listStrContent[7], wsWordData)
        wsWord.merge_range(intTestProfileStartLine + 4, 1, intTestProfileStartLine + 4, intActualMeasuredCapacityLength, listStrContent[8], wsWordData)
        wsWord.merge_range(intTestProfileStartLine + 5, 1, intTestProfileStartLine + 5, intActualMeasuredCapacityLength, listStrContent[9], wsWordData)
        wsWord.merge_range(intTestProfileStartLine + 6, 1, intTestProfileStartLine + 6, intActualMeasuredCapacityLength, listStrContent[10], wsWordData)
        wsWord.merge_range(intTestProfileStartLine + 7, 1, intTestProfileStartLine + 7, int(intActualMeasuredCapacityLength / 2), listStrContent[11], wsWordData)
        wsWord.merge_range(intTestProfileStartLine + 7, int(intActualMeasuredCapacityLength / 2) + 1, intTestProfileStartLine + 7, intActualMeasuredCapacityLength, listStrContent[12], wsWordData)
        for v in range(self.intVoltageLevelNum):
            if v == intPosi2V25:
                wsWord.merge_range(intTestProfileStartLine + 8, 1 + v * 2, intTestProfileStartLine + 8, 2 + v * 2, f"{self.listVoltageLevel[v]}V", wsWordData_bold)
                wsWord.merge_range(intTestProfileStartLine + 9, 1 + v * 2, intTestProfileStartLine + 9, 2 + v * 2, f"{round(listMM2S[intPosiMaxmA][v], 2)}", wsWordData_bold)
                wsWord.merge_range(intTestProfileStartLine + 10, 1 + v * 2, intTestProfileStartLine + 10, 2 + v * 2, "", wsWordData_bold)
                wsWord.write_formula(f"{excel_utils.num2letter(1 + v * 2)}{intTestProfileStartLine + 11}",
                                     f"=TRUNC({excel_utils.num2letter(1 + v * 2)}{intTestProfileStartLine + 10}/B{intTestProfileStartLine + 7}, 2)",
                                     wsWordData_percentage_bold)
            else:
                wsWord.merge_range(intTestProfileStartLine + 8, 1 + v * 2, intTestProfileStartLine + 8, 2 + v * 2, f"{self.listVoltageLevel[v]}V", wsWordData)
                wsWord.merge_range(intTestProfileStartLine + 9, 1 + v * 2, intTestProfileStartLine + 9, 2 + v * 2, f"{round(listMM2S[intPosiMaxmA][v], 2)}", wsWordData)
                wsWord.merge_range(intTestProfileStartLine + 10, 1 + v * 2, intTestProfileStartLine + 10, 2 + v * 2, "", wsWordData)
                wsWord.write_formula(f"{excel_utils.num2letter(1 + v * 2)}{intTestProfileStartLine + 11}",
                                     f"=TRUNC({excel_utils.num2letter(1 + v * 2)}{intTestProfileStartLine + 10}/B{intTestProfileStartLine + 7}, 2)",
                                     wsWordData_percentage)

        wsWord.merge_range(intTestDateStartRow, 1, intTestDateStartRow, intActualMeasuredCapacityLength, listStrContent[14], wsWordData)
        wsWord.merge_range(intTestDateStartRow + 1, 1, intTestDateStartRow + 1, intActualMeasuredCapacityLength, listStrContent[15], wsWordData)
        wsWord.merge_range(intTestDateStartRow + 2, 1, intTestDateStartRow + 2, intActualMeasuredCapacityLength, listStrContent[16], wsWordData)
        wsWord.merge_range(intTestDateStartRow + 3, 1, intTestDateStartRow + 3, intActualMeasuredCapacityLength, listStrContent[17], wsWordData_bgyellow)
        wsWord.merge_range(intTestDateStartRow + 4, 1, intTestDateStartRow + 4, intActualMeasuredCapacityLength, "", wsWordData)
        # 将相对路径转换为file:// URL格式
        url_path = strRelResultPath
        # 确保路径使用正斜杠
        url_path = url_path.replace('\\', '/')
        # 添加file://前缀
        file_url = f'file:///{url_path}'
        wsWord.write_url(intTestDateStartRow + 4, 1, file_url, wbSampleHyperlink, string=listStrContent[18].split("\\")[-1])
        wsWord.merge_range(intTestDateStartRow + 5, 1, intTestDateStartRow + 5, intActualMeasuredCapacityLength, listStrContent[19], wsWordData)

        # wdResult write table Overview
        tableOverview = wdReport.add_table(intTestDateStartRow + 6, 1 + self.intVoltageLevelNum * 2, style='Table Grid')
        for row in range(intTestDateStartRow + 6):
            tableOverview.rows[row].height = Cm(0.5)
            for col in range(1 + self.intVoltageLevelNum * 2):
                cell = tableOverview.cell(row, col)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.line_spacing_rules = WD_LINE_SPACING.SINGLE
                cell.paragraphs[0].paragraph_format.space_after = Pt(0)
                if row == intTestDateStartRow + 3:
                    word_utils.table_set_bg_color(cell, '#FFFF00')

        # merge cell
        for row in range(intTestDateStartRow + 6):
            if row <= intTestProfileStartLine + 6 or row >= intTestDateStartRow:
                cell1 = tableOverview.cell(row, 1)
                for col in range(1, self.intVoltageLevelNum * 2):
                    cell2 = tableOverview.cell(row, col + 1)
                    cell1.merge(cell2)
            else:
                if row == intTestProfileStartLine + 7:
                    cell1 = tableOverview.cell(row, 1)
                    for col in range(1, self.intVoltageLevelNum):
                        cell2 = tableOverview.cell(row, col + 1)
                        cell1.merge(cell2)
                    cell1 = tableOverview.cell(row, 1 + self.intVoltageLevelNum)
                    for col in range(1 + self.intVoltageLevelNum, self.intVoltageLevelNum * 2):
                        cell2 = tableOverview.cell(row, col + 1)
                        cell1.merge(cell2)
                else:
                    for v in range(self.intVoltageLevelNum):
                        cell1 = tableOverview.cell(row, 1 + v * 2)
                        cell2 = tableOverview.cell(row, 2 + v * 2)
                        cell1.merge(cell2)

        cell1 = tableOverview.cell(intTestProfileStartLine + 8, 0)
        cell2 = tableOverview.cell(intTestProfileStartLine + 9, 0)
        cell1.merge(cell2)
        cell2 = tableOverview.cell(intTestProfileStartLine + 10, 0)
        cell1.merge(cell2)

        tableOverview.cell(0, 0).paragraphs[0].add_run(listStrItems[0])
        tableOverview.cell(1, 0).paragraphs[0].add_run(listStrItems[1])
        tableOverview.cell(2, 0).paragraphs[0].add_run(listStrItems[2])
        tableOverview.cell(3, 0).paragraphs[0].add_run(listStrItems[3])
        if intTestProfileStartLine == 4:
            tableOverview.cell(3, 0).paragraphs[0].text = tableOverview.cell(3, 0).paragraphs[0].text.replace(listStrItems[3], "")
        tableOverview.cell(intTestProfileStartLine, 0).paragraphs[0].add_run(listStrItems[4])
        tableOverview.cell(intTestProfileStartLine + 1, 0).paragraphs[0].add_run(listStrItems[5])
        tableOverview.cell(intTestProfileStartLine + 2, 0).paragraphs[0].add_run(listStrItems[6])
        tableOverview.cell(intTestProfileStartLine + 3, 0).paragraphs[0].add_run(listStrItems[7])
        tableOverview.cell(intTestProfileStartLine + 4, 0).paragraphs[0].add_run(listStrItems[8])
        tableOverview.cell(intTestProfileStartLine + 5, 0).paragraphs[0].add_run(listStrItems[9])
        tableOverview.cell(intTestProfileStartLine + 6, 0).paragraphs[0].add_run(listStrItems[10])
        tableOverview.cell(intTestProfileStartLine + 7, 0).paragraphs[0].add_run(listStrItems[11])
        tableOverview.cell(intTestProfileStartLine + 8, 0).paragraphs[0].add_run(listStrItems[13])
        tableOverview.cell(intTestDateStartRow, 0).paragraphs[0].add_run(listStrItems[14])
        tableOverview.cell(intTestDateStartRow + 1, 0).paragraphs[0].add_run(listStrItems[15])
        tableOverview.cell(intTestDateStartRow + 2, 0).paragraphs[0].add_run(listStrItems[16])
        tableOverview.cell(intTestDateStartRow + 3, 0).paragraphs[0].add_run(listStrItems[17])
        tableOverview.cell(intTestDateStartRow + 4, 0).paragraphs[0].add_run(listStrItems[18])
        tableOverview.cell(intTestDateStartRow + 5, 0).paragraphs[0].add_run(listStrItems[19])

        tableOverview.cell(0, 1).paragraphs[0].add_run(listStrContent[0])
        tableOverview.cell(1, 1).paragraphs[0].add_run(listStrContent[1])
        tableOverview.cell(2, 1).paragraphs[0].add_run(listStrContent[2])
        tableOverview.cell(3, 1).paragraphs[0].add_run(listStrContent[3])
        if intTestProfileStartLine == 4:
            tableOverview.cell(3, 1).paragraphs[0].text = tableOverview.cell(3, 1).paragraphs[0].text.replace(listStrContent[3], "")
        tableOverview.cell(intTestProfileStartLine, 1).paragraphs[0].text = ""
        if len(listStrContent[4].split("\\")) == 1:
            tableOverview.cell(intTestProfileStartLine, 1).paragraphs[0].add_run(listStrContent[4])
        else:
            word_utils.add_hyperlink(tableOverview.cell(intTestProfileStartLine, 1).paragraphs[0], listStrContent[4], listStrContent[4].split("\\")[-1])
        tableOverview.cell(intTestProfileStartLine + 1, 1).paragraphs[0].add_run(listStrContent[5])
        tableOverview.cell(intTestProfileStartLine + 2, 1).paragraphs[0].add_run(listStrContent[6])
        tableOverview.cell(intTestProfileStartLine + 3, 1).paragraphs[0].add_run(listStrContent[7])
        tableOverview.cell(intTestProfileStartLine + 4, 1).paragraphs[0].add_run(listStrContent[8])
        tableOverview.cell(intTestProfileStartLine + 5, 1).paragraphs[0].add_run(listStrContent[9])
        tableOverview.cell(intTestProfileStartLine + 6, 1).paragraphs[0].add_run(listStrContent[10])
        tableOverview.cell(intTestProfileStartLine + 7, 1).paragraphs[0].add_run(listStrContent[11])
        tableOverview.cell(intTestProfileStartLine + 7, 1 + self.intVoltageLevelNum).paragraphs[0].add_run(listStrContent[12])
        for v in range(self.intVoltageLevelNum):
            text1 = tableOverview.cell(intTestProfileStartLine + 8, 1 + v * 2).paragraphs[0].add_run(f"{self.listVoltageLevel[v]}V")
            text2 = tableOverview.cell(intTestProfileStartLine + 9, 1 + v * 2).paragraphs[0].add_run(f"{round(listMM2S[intPosiMaxmA][v], 2)}")
            text3 = tableOverview.cell(intTestProfileStartLine + 10, 1 + v * 2).paragraphs[0].add_run(f"{math.floor(100 * listMM2S[intPosiMaxmA][v] / int(listStrContent[10]))}%")
            if v == intPosi2V25:
                text1.font.bold = True
                text2.font.bold = True
                text3.font.bold = True
        tableOverview.cell(intTestDateStartRow, 1).paragraphs[0].add_run(listStrContent[14])
        tableOverview.cell(intTestDateStartRow + 1, 1).paragraphs[0].add_run(listStrContent[15])
        tableOverview.cell(intTestDateStartRow + 2, 1).paragraphs[0].add_run(listStrContent[16])
        tableOverview.cell(intTestDateStartRow + 3, 1).paragraphs[0].add_run(listStrContent[17])
        tableOverview.cell(intTestDateStartRow + 4, 1).paragraphs[0].text = ""
        word_utils.add_hyperlink(tableOverview.cell(intTestDateStartRow + 4, 1).paragraphs[0], listStrContent[18], listStrContent[18].split("\\")[-1])
        tableOverview.cell(intTestDateStartRow + 5, 1).paragraphs[0].add_run(listStrContent[19])

        for row in range(intTestDateStartRow + 6):
            for col in range(1 + self.intVoltageLevelNum * 2):
                cell = tableOverview.cell(row, col)
                runs = cell.paragraphs[0].runs
                for run in runs:
                    run.font.size = Pt(9)
        tableOverview.cell(0, 0).width = Cm(27)

        # 合并三次文档遍历时为一次，减少I/O操作
        # wdResult replace TypeA-TypeF, StrA-StrF, Image and insert tables in one pass
        bInsertOverview = False
        bInsertVersionHistory = False
        bInsertTestInformation = False
        bInsertStatisticalsResults = False
        intStepOut = 0
        
        for paragraph in wdReport.paragraphs:
            # 1. 替换文本和插入图片
            modified = False
            for t in range(len(self.listTextToReplace)):
                if self.listTextToReplace[t] in paragraph.text:
                    modified = True
                    if self.listTextToReplace[t] == "StrD":
                        paragraph.text = paragraph.text.replace(self.listTextToReplace[t], "")
                        text = paragraph.add_run(f"{self.listTestInfoForReplace[t]}")
                        text.font.bold = True
                        paragraph.add_run(".")
                    else:
                        paragraph.text = paragraph.text.replace(self.listTextToReplace[t], f"{self.listTestInfoForReplace[t]}")
            
            if not modified:  # 只有当段落未被修改时才处理图片，避免重复处理
                for i in range(len(self.listImageToReplace)):
                    if self.listImageToReplace[i] in paragraph.text:
                        paragraph.text = paragraph.text.replace(self.listImageToReplace[i], "")
                        if i < 2 * len(self.listPngPath) + 1:
                            if i == 0:
                                paragraph.add_run("").add_picture(self.strFilteredPngPath, width=Cm(15))
                            elif i % 2 == 1:
                                paragraph.add_run("").add_picture(self.listPngPath[int((i - 1) / 2)], width=Cm(7.2))
                            else:
                                paragraph.add_run(f"Figure {int(i / 2 + 1)}  {self.listTestInfoForReplace[2]} {self.listTestInfoForReplace[0]}-{self.listTestInfoForReplace[1]} Boxplot, {self.listCurrentLevel[int(i / 2 - 1)]}mA")
                        else:
                            paragraph._element.getparent().remove(paragraph._element)
                            continue  # 跳过后续处理，因为段落已被删除
            
            # 2. 插入表格的逻辑
            if "Battery Quality Test / Alternative Battery Test for ESL Batteries" in paragraph.text:
                bInsertOverview = True
                intStepOut = 4
            elif "Version history" in paragraph.text and "Heading 2" == paragraph.style.name:
                bInsertVersionHistory = True
                intStepOut = 0
            elif "Test Information" in paragraph.text and "Heading 1" == paragraph.style.name:
                bInsertTestInformation = True
                intStepOut = 0
            elif "Test results" in paragraph.text and "Heading 1" == paragraph.style.name:
                bInsertStatisticalsResults = True
                intStepOut = 2
            
            if intStepOut:
                intStepOut = intStepOut - 1
            else:
                if bInsertOverview:
                    bInsertOverview = False
                    paragraph._p.addnext(tableOverview._tbl)
                elif bInsertVersionHistory:
                    bInsertVersionHistory = False
                    paragraph._p.addnext(tableVersionHistory._tbl)
                elif bInsertTestInformation:
                    bInsertTestInformation = False
                    paragraph._p.addnext(tableTestInformation._tbl)
                elif bInsertStatisticalsResults:
                    bInsertStatisticalsResults = False
                    paragraph._p.addnext(tableStatisticalsResults._tbl)
            
            # 3. 删除温度符号
            if listStrContent[16] == "Room Temperature" and "\u2103" in paragraph.text:
                paragraph.text = paragraph.text.replace("\u2103", "")

        # close xlsx writer
        wbResult.close()
        wbSample.close()
        # close word writer
        wdReport.save(self.strReportWordPath)
        # 输出docx文件的完整路径到日志
        logging.info(f"数据分析完成，生成的docx报告路径: {self.strReportWordPath}")
        # close csv writer - 确保缓冲区中的所有数据都被写入
        if csv_buffer:
            csvwriterResultCsvFile.writerows(csv_buffer)
            csv_buffer.clear()
        f.close()

    def UXWW_Draw(self, _listCpt: list, maxXaxis: int) -> None:
        fontdictLabel = {
            'fontsize': 9,
            'fontweight': 'bold'
        }
        medianprofile = dict(linewidth=1, color='red')

        plt.figure()

        for c in range(self.intCurrentLevelNum):
            listBoxPlot = []
            listLabel = []
            for v in range(self.intVoltageLevelNum):
                listBoxPlot.append(_listCpt[c][v])
                listLabel.append(f"{self.listVoltageLevel[v]}V")
            plt.cla()
            plt.boxplot(listBoxPlot, labels=listLabel, medianprops=medianprofile)
            plt.title(self.listBoxplotTitle[c], fontdict=fontdictLabel)
            plt.xlabel("Cutoff Voltage [V]")
            plt.ylabel("Useable Capacity [mAh]")
            plt.grid(linestyle="--", alpha=0.3)
            plt.savefig(self.listPngPath[c])
            plt.savefig(self.listSvgPath[c], dpi=1200)

        # analysis Info_Image.csv
        listPlt = []
        for c in range(self.intCurrentLevelNum):
            listPlt.append([])
            for _ in range(4):
                listPlt[c].append([])

        f = open(self.strInfoImageCsvPath, mode='r', encoding='utf-8')
        csvreaderInfoImageCsvFile = csv.reader(f)
        intPerBatteryRows = 1 + self.intCurrentLevelNum * 3
        index = 0
        for row in csvreaderInfoImageCsvFile:
            loop = index % intPerBatteryRows
            if loop != 0 and (loop % 3) != 1:
                listPlt[int((loop - 1) / 3)][((loop - 1) % 3) - 1].append([float(row[i]) for i in range(len(row))])
            index += 1
        f.close()

        for c in range(self.intCurrentLevelNum):
            listPlt[c][2], listPlt[c][3] = data_utils.filter_data(listPlt[c][0], listPlt[c][1])

        title_fontdict = {
            'fontsize': 15,
            'fontweight': 'bold'
        }
        axis_fontdict = {
            'fontsize': 15
        }

        plt.figure(figsize=(15, 6))

        plt.clf()
        plot_utils.set_plt_axis(self.listTestInfo[0], maxXaxis)
        y_major_locator = MultipleLocator(0.2)
        ax = plt.gca()
        ax.yaxis.set_major_locator(y_major_locator)
        plt.title(f"Unfiltered {self.strPltName}", fontdict=title_fontdict)
        plt.xlabel("Charge [mAh]", fontdict=axis_fontdict)
        plt.ylabel("Unfiltered Battery Load Voltage [V]", fontdict=axis_fontdict)
        for b in range(self.intBatteryNum):
            for c in range(self.intCurrentLevelNum):
                plt.plot(listPlt[c][0][b], listPlt[c][1][b], color=f"{self.listPltColorType[c]}", linewidth=0.5)
        plt.grid(linestyle="--", alpha=0.3)
        plt.savefig(self.strUnfilteredPngPath)
        plt.savefig(self.strUnfilteredSvgPath, dpi=1200)

        plt.clf()
        plot_utils.set_plt_axis(self.listTestInfo[0], maxXaxis)
        y_major_locator = MultipleLocator(0.2)
        ax = plt.gca()
        ax.yaxis.set_major_locator(y_major_locator)
        plt.title(f"Filtered {self.strPltName}", fontdict=title_fontdict)
        plt.xlabel("Charge [mAh]", fontdict=axis_fontdict)
        plt.ylabel("Filtered Battery Load Voltage [V]", fontdict=axis_fontdict)
        for b in range(self.intBatteryNum):
            for c in range(self.intCurrentLevelNum):
                plt.plot(listPlt[c][0][b], listPlt[c][1][b], color=f"{self.listPltColorType[c]}", linewidth=0.5)
        plt.grid(linestyle="--", alpha=0.3)
        plt.savefig(self.strFilteredPngPath)
        plt.savefig(self.strFilteredSvgPath, dpi=1200)


class JsonWriter:
    def __init__(self, strResultPath: str, listTestInfo: list, listBatteryInfo: list) -> None:
        self.config = configparser.ConfigParser()
        
        try:
            # 使用通用配置文件查找函数
            config_path = find_config_file()
            if config_path and os.path.exists(config_path):
                self.config.read(config_path, encoding='utf-8')
                logging.info(f"找到并读取配置文件: {config_path}")
            else:
                raise Exception("找不到配置文件")
        except Exception as e:
            # 发生错误时创建基本配置
            logging.error(f"配置读取失败: {e}，使用默认配置")
            if not self.config.has_section("BatteryConfig"):
                self.config.add_section("BatteryConfig")
            if not self.config.has_section("PltConfig"):
                self.config.add_section("PltConfig")
        self.listTestInfo = listTestInfo
        self.listBatteryInfo = listBatteryInfo
        try:
            [sy, sm, sd] = self.listBatteryInfo[2][0].split(" ")[0].split("-")
            td = f"{sy}{sm}{sd}"
        except ValueError:
            td = "00000000"
        self.strResultPath = f"{strResultPath}/{td}_V{listTestInfo[16]}"
        self.dictJson = {}
        self.listTestRun = []
        self.dictMeasurements = {}
        self.runAt = datetime.datetime.strptime(self.listBatteryInfo[2][0], "%Y-%m-%d %H:%M:%S").astimezone(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")

        self.listCurrentLevel = listTestInfo[14]
        self.listVoltageLevel = listTestInfo[15]
        self.strFileCurrentType = ""
        for c in range(len(self.listCurrentLevel)):
            self.strFileCurrentType = self.strFileCurrentType + f"{self.listCurrentLevel[c]}-"
        self.strFileCurrentType = self.strFileCurrentType[:-1]
        # 使用os.path.join确保路径分隔符一致性
        self.strResultJsonPath = os.path.join(self.strResultPath, f"{self.listTestInfo[4]}_{self.listTestInfo[2]}_{self.listTestInfo[3]}_{self.strFileCurrentType}_{self.listTestInfo[7]}.json")
        self.listBatteryVoltage = []
        for v in range(len(self.listVoltageLevel)):
            self.listBatteryVoltage.append(len(str(self.listVoltageLevel[v])) < 4 and str(self.listVoltageLevel[v]) + '0' * (4 - len(str(self.listVoltageLevel[v]))) or str(self.listVoltageLevel[v]))

        self.UJS_FormatJson()

    def UJS_FormatJson(self) -> None:
        for i in range(len(self.listBatteryInfo[1])):
            index = -1
            dictMeasurements = []
            for c in range(len(self.listCurrentLevel)):
                dictMeasurements.append({})
            dictTestRun = {}
            for j in range(len(self.listBatteryInfo[0][i])):
                if j % (len(self.listVoltageLevel)) == 0:
                    index += 1
                if self.listBatteryInfo[0][i][j] != 0 and self.listBatteryVoltage[j % len(self.listBatteryVoltage)] != "":
                    dictMeasurements[index].update({self.listBatteryVoltage[j % len(self.listBatteryVoltage)]: self.listBatteryInfo[0][i][j]})
            try:
                battery_name_split = self.listBatteryInfo[1][i].split("BTS")[1].split("_")
                battery_name = f"{battery_name_split[2]}_{battery_name_split[3]}"
            except IndexError:
                battery_name = f"Battery_{i}"
            listResults = []
            for c in range(len(self.listCurrentLevel)):
                listResults.append({"scenario": f"{self.listCurrentLevel[c]}mA", "measurements": dictMeasurements[c]})
            dictTestRun.update({"slot": battery_name, "results": listResults})
            self.listTestRun.append(dictTestRun)

        strBatteryModel = self.listTestInfo[2]
        # 安全读取电池类型基础规格
        try:
            if self.config.has_section("BatteryConfig") and self.config.has_option("BatteryConfig", "SpecificationTypeBase"):
                listBatteryTypeBase = self.config.get("BatteryConfig", "SpecificationTypeBase").split(",")
                logging.info(f"使用配置文件中的电池类型基础规格: {listBatteryTypeBase}")
            else:
                # 使用默认值
                listBatteryTypeBase = ["CoinCell", "ButtonCell", "Cylindrical", "Prismatic", "PouchCell"]
                logging.info("使用默认电池类型基础规格")
            
            strBatteryType = ""
            for b in range(len(listBatteryTypeBase)):
                if listBatteryTypeBase[b].strip() in self.listTestInfo[2]:
                    strBatteryType = listBatteryTypeBase[b]
                    break
            
            # 如果没有找到匹配项，使用默认值
            if strBatteryType == "":
                strBatteryType = "CoinCell"
                logging.warning(f"未找到精确匹配的电池类型，使用默认值: {strBatteryType}")
        except Exception as e:
            logging.error(f"处理电池类型时出错: {e}，使用默认值")
            strBatteryType = "CoinCell"

        self.dictJson.update({
            "batchId": self.listTestInfo[5],
            "runAt": self.runAt,
            "batteryType": strBatteryType,
            "batteryModel": strBatteryModel,
            "batteryManufacturer": self.listTestInfo[4],
            "testRuns": self.listTestRun})

        # 确保目标目录存在
        os.makedirs(os.path.dirname(self.strResultJsonPath), exist_ok=True)
        with open(self.strResultJsonPath, 'w') as file:
            json.dump(self.dictJson, file, indent=4)


class FileWriter:
    def __init__(self, strResultPath: str, listTestInfo: list, listBatteryInfo: list) -> None:
        self.strErrorLog = ""
        try:
            XlsxWordWriter(strResultPath, listTestInfo, listBatteryInfo)
            JsonWriter(strResultPath, listTestInfo, listBatteryInfo)
        except BaseException as e:
            self.strErrorLog = str(e)
            traceback.print_exc()

    def UFW_GetErrorLog(self) -> str:
        return self.strErrorLog
